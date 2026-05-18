from docx import Document
from docx.shared import Inches
from io import BytesIO
import os
from agents.consultation.consultation_agent import SessionLocal
from agents.specification.specification_agent import GeneratedSpecification
from agents.claim.claim_agent import GeneratedClaim
from agents.drawing.drawing_db import GeneratedDrawing


def get_patent_data(user_id, consultation_idx):
    db = SessionLocal()
    try:
        spec = db.query(GeneratedSpecification).filter_by(
            user_id=user_id, consultation_idx=consultation_idx
        ).first()

        claim = db.query(GeneratedClaim).filter_by(
            user_id=user_id, consultation_idx=consultation_idx
        ).first()

        drawings = db.query(GeneratedDrawing).filter_by(
            user_id=user_id, consultation_idx=consultation_idx
        ).order_by(GeneratedDrawing.fig_number).all()

        return spec, claim, drawings
    finally:
        db.close()


def generate_patent_docx(user_id, consultation_idx):
    spec, claim, drawings = get_patent_data(user_id, consultation_idx)

    doc = Document()

    doc.add_heading('특허 명세서', 0)

    doc.add_paragraph('【발명의 명칭】')
    doc.add_paragraph('인공지능 기반 자동 명세서 작성 시스템')

    doc.add_paragraph('【기술분야】')
    doc.add_paragraph(spec.tech_field if spec else "내용 없음")

    doc.add_paragraph('【배경기술】')
    doc.add_paragraph(spec.background_art if spec else "내용 없음")

    doc.add_paragraph('【발명의 내용】')
    doc.add_paragraph('【해결하려는 과제】')
    doc.add_paragraph(spec.problem_statement if spec else "내용 없음")

    doc.add_paragraph('【과제의 해결 수단】')
    doc.add_paragraph(spec.solution_means if spec else "내용 없음")

    doc.add_paragraph('【발명의 효과】')
    doc.add_paragraph(spec.effects if spec else "내용 없음")

    doc.add_paragraph('【도면의 간단한 설명】')
    drawing_desc = spec.drawing_description if spec else None
    if not drawing_desc and drawings:
        auto_desc = "\n".join([
            f"【{d.fig_number}】은 본 발명의 일 실시예에 따른 도면이다."
            for d in drawings
        ])
        doc.add_paragraph(auto_desc)
    else:
        doc.add_paragraph(drawing_desc or "도면 없음")

    doc.add_paragraph('【발명을 실시하기 위한 구체적인 내용】')
    detailed_text = (spec.detailed_desc or "") if spec else ""
    embodiments_text = (spec.embodiments or "") if spec else ""
    combined_detail = f"{detailed_text}\n\n{embodiments_text}".strip()
    doc.add_paragraph(combined_detail or "내용 없음")

    doc.add_page_break()
    doc.add_heading('특허청구범위', level=1)

    doc.add_paragraph('【청구항 1】')
    doc.add_paragraph(claim.claim_1 if claim else "내용 없음")

    if claim and claim.dependent_claims:
        doc.add_paragraph(claim.dependent_claims)

    if drawings:
        doc.add_page_break()
        doc.add_heading('도면', level=1)
        for d in drawings:
            doc.add_paragraph(f"【{d.fig_number}】")
            if d.png_path and os.path.exists(d.png_path):
                doc.add_picture(d.png_path, width=Inches(5.5))
            else:
                doc.add_paragraph("[도면 이미지 파일을 찾을 수 없습니다]")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
