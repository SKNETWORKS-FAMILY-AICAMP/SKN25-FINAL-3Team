import os
import re
from typing import Any, Dict, List, Optional
from docx import Document
from docx.shared import Inches
from langchain_openai import ChatOpenAI
from langchain_core.prompts import ChatPromptTemplate

def validate_composer_inputs(state: Dict[str, Any]) -> None:
    # Optional preliminary check. The adapter functions already raise errors if critical parts are missing.
    pass

def get_claim_1_text(state: Dict[str, Any]) -> str:
    if state.get("claim_1_text"):
        return state["claim_1_text"]
    
    claims_obj = state.get("claims")
    if isinstance(claims_obj, list) and len(claims_obj) > 0:
        first_claim = claims_obj[0]
        if isinstance(first_claim, dict):
            return first_claim.get("text", "")
        return str(first_claim)
    elif isinstance(claims_obj, dict):
        draft_claims = claims_obj.get("draft_claims", [])
        if draft_claims and len(draft_claims) > 0:
            first_claim = draft_claims[0]
            if isinstance(first_claim, dict):
                return first_claim.get("text", "")
            return str(first_claim)
            
    if state.get("claims_text"):
        match = re.search(r"【청구항 1】(.*?)(?:【청구항 2】|$)", state["claims_text"], re.DOTALL)
        if match:
            return match.group(1).strip()
        return state["claims_text"].strip()
        
    raise ValueError("청구항 1항을 찾을 수 없습니다.")

def get_all_claims_text(state: Dict[str, Any]) -> str:
    if state.get("claims_text"):
        return state["claims_text"]
        
    claims_obj = state.get("claims")
    claims_list = []
    if isinstance(claims_obj, list):
        claims_list = claims_obj
    elif isinstance(claims_obj, dict):
        claims_list = claims_obj.get("draft_claims", [])
        
    if claims_list:
        texts = []
        for i, c in enumerate(claims_list):
            if isinstance(c, dict):
                text = c.get("text", "")
                claim_no = c.get("claim_no", i + 1)
                texts.append(f"【청구항 {claim_no}】\n{text}")
            else:
                texts.append(f"【청구항 {i+1}】\n{c}")
        return "\n\n".join(texts)
        
    raise ValueError("청구항 전체 내용을 찾을 수 없습니다.")

def get_drawings(state: Dict[str, Any]) -> List[Dict[str, Any]]:
    drawings_obj = state.get("drawings")
    if isinstance(drawings_obj, list):
        return drawings_obj
    elif isinstance(drawings_obj, dict) and "figures" in drawings_obj:
        return drawings_obj["figures"]
    return []

def get_drawing_image_path(drawing: Dict[str, Any]) -> Optional[str]:
    # Check in specific order: image_path -> png_path -> svg_path
    for key in ["image_path", "png_path", "svg_path"]:
        if drawing.get(key):
            return drawing[key]
    return None

def get_specification_sections(state: Dict[str, Any]) -> Dict[str, str]:
    spec_obj = state.get("specification_sections")
    if spec_obj and isinstance(spec_obj, dict):
        return spec_obj
    
    spec_obj = state.get("specification")
    if spec_obj and isinstance(spec_obj, dict):
        return spec_obj
        
    return {}

def generate_abstract_from_claim_1(claim_1_text: str) -> str:
    model_name = os.environ.get("COMPOSER_MODEL", "gpt-4o")
    llm = ChatOpenAI(model=model_name, temperature=0.2)
    
    prompt = ChatPromptTemplate.from_messages([
        ("system", "너는 특허 명세서 요약문 작성 보조자이다.\n"
                   "입력된 청구항 1항을 바탕으로 【요약】 항목에 들어갈 문장을 작성하라.\n"
                   "청구항 1항의 기술적 구성을 유지하라.\n"
                   "없는 구성요소를 새로 만들지 마라.\n"
                   "없는 효과를 새로 만들지 마라.\n"
                   "권리범위를 불필요하게 축소하지 마라.\n"
                   "\"본 발명은 ...에 관한 것으로서\" 문체를 사용하라.\n"
                   "1문단 또는 2문단으로 작성하라.\n"
                   "특허 명세서 요약문에 어울리는 문체로 작성하라."),
        ("user", "청구항 1항:\n{claim_1_text}")
    ])
    
    chain = prompt | llm
    response = chain.invoke({"claim_1_text": claim_1_text})
    return response.content

def select_representative_drawing(state: Dict[str, Any]) -> Optional[str]:
    drawings = get_drawings(state)
    if drawings and len(drawings) > 0:
        return get_drawing_image_path(drawings[0])
    
    if state.get("drawing_image_paths") and len(state["drawing_image_paths"]) > 0:
        return state["drawing_image_paths"][0]
        
    return None

def create_final_docx(
    output_path: str,
    abstract_text: str,
    representative_drawing_path: Optional[str],
    claims_text: str,
    spec_sections: Dict[str, str],
    drawings: List[Dict[str, Any]]
):
    doc = Document()
    
    # 1. 요약
    doc.add_heading("요약", level=1)
    doc.add_paragraph(abstract_text)
    doc.add_page_break()
    
    # 2. 대표도
    doc.add_heading("대표도", level=1)
    if representative_drawing_path and os.path.exists(representative_drawing_path):
        try:
            doc.add_picture(representative_drawing_path, width=Inches(6.0))
        except Exception as e:
            doc.add_paragraph(f"[대표도 이미지 삽입 실패: {e}]")
    else:
        doc.add_paragraph("대표도 없음")
    doc.add_page_break()
    
    # 3. 청구항/청구범위
    doc.add_heading("청구항", level=1)
    doc.add_paragraph("【청구범위】")
    doc.add_paragraph(claims_text)
    doc.add_page_break()
    
    # 4. 발명의 설명
    doc.add_heading("발명의 설명", level=1)
    
    # 【발명(고안)의 설명】 or 【발명(고안)의 명칭】 MUST NOT BE INCLUDED
    # Directly start with 【기술분야】
    
    def add_spec_section(title, key1, key2=None):
        content = spec_sections.get(key1, "")
        if not content and key2:
            content = spec_sections.get(key2, "")
        if content:
            doc.add_paragraph(title)
            doc.add_paragraph(content)

    add_spec_section("【기술분야】", "technical_field", "기술분야")
    add_spec_section("【발명(고안)의 배경이 되는 기술】", "background_art", "배경기술")
    
    # 【발명(고안)의 내용】 serves as a grouping header
    doc.add_paragraph("【발명(고안)의 내용】")
    add_spec_section("【해결하려는 과제】", "problem_to_solve", "problem")
    add_spec_section("【과제의 해결 수단】", "means_for_solving", "solution")
    add_spec_section("【발명(고안)의 효과】", "effects", "effect")
    
    add_spec_section("【도면의 간단한 설명】", "brief_description_of_drawings", "도면의 간단한 설명")
    add_spec_section("【발명(고안)을 실시하기 위한 구체적인 내용】", "detailed_description", "구체적인 실시예")
    doc.add_page_break()
    
    # 5. 도면
    doc.add_heading("도면", level=1)
    for i, drawing in enumerate(drawings):
        fig_no = drawing.get("figure_no", drawing.get("fig_no", i + 1))
        if isinstance(fig_no, int) or str(fig_no).isdigit():
            fig_str = f"【도 {fig_no}】"
        else:
            if str(fig_no).startswith("도 "):
                fig_str = f"【{fig_no}】"
            else:
                fig_str = f"【도 {fig_no}】"
                
        doc.add_paragraph(fig_str)
        
        desc = drawing.get("description", "")
        if desc:
            doc.add_paragraph(desc)
            
        img_path = get_drawing_image_path(drawing)
        if img_path and os.path.exists(img_path):
            try:
                doc.add_picture(img_path, width=Inches(6.0))
            except Exception as e:
                doc.add_paragraph(f"[도면 이미지 삽입 실패: {e}]")
        else:
            doc.add_paragraph("[도면 이미지 없음]")
            
        doc.add_paragraph("")
        
    doc.save(output_path)

def run_composer_agent(state: Dict[str, Any]) -> Dict[str, Any]:
    # 1. Extract and validate
    claim_1_text = get_claim_1_text(state)
    all_claims_text = get_all_claims_text(state)
    spec_sections = get_specification_sections(state)
    drawings = get_drawings(state)
    
    # 2. Generate abstract
    abstract_text = generate_abstract_from_claim_1(claim_1_text)
    
    # 3. Select representative drawing
    rep_drawing_path = select_representative_drawing(state)
    
    # 4. Create Word Document
    os.makedirs("outputs", exist_ok=True)
    final_docx_path = os.path.abspath("outputs/final_patent_specification.docx")
    
    create_final_docx(
        output_path=final_docx_path,
        abstract_text=abstract_text,
        representative_drawing_path=rep_drawing_path,
        claims_text=all_claims_text,
        spec_sections=spec_sections,
        drawings=drawings
    )
    
    # 5. Update state
    if "final_package" not in state:
        state["final_package"] = {}
        
    state["final_package"]["rendered_docx_path"] = final_docx_path
    state["final_docx_path"] = final_docx_path
    state["abstract_text"] = abstract_text
    state["representative_drawing_path"] = rep_drawing_path
    
    return state
