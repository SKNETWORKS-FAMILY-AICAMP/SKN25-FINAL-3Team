import os
from datetime import datetime
from typing import Any, Dict, List, Optional

from docx import Document
from docx.shared import Inches

from .adapters import get_drawing_image_path, sanitize_filename

SECTION_MAPPING = [
    ("【기술분야】", ["technical_field", "기술분야"]),
    ("【발명(고안)의 배경이 되는 기술】", ["background_art", "background", "배경기술"]),
    ("【해결하려는 과제】", ["problem_to_solve", "problem", "summary_problem", "해결하려는 과제"]),
    (
        "【과제의 해결 수단】",
        ["means_for_solving", "solution", "summary_solution", "과제의 해결수단", "과제의 해결 수단"],
    ),
    ("【발명(고안)의 효과】", ["effects", "effect", "summary_effect", "발명의 효과"]),
    ("【도면의 간단한 설명】", ["brief_description_of_drawings", "brief_drawings", "도면의 간단한 설명"]),
    (
        "【발명(고안)을 실시하기 위한 구체적인 내용】",
        ["detailed_description", "embodiment", "구체적인 실시예", "구체적인 내용"],
    ),
]


def is_supported_docx_image(path: str) -> bool:
    return os.path.splitext(path)[1].lower() in [".png", ".jpg", ".jpeg", ".bmp", ".gif", ".tif", ".tiff"]


def add_image_or_placeholder(doc: Document, image_path: Optional[str], width_inches: float = 6.0) -> None:
    if not image_path:
        doc.add_paragraph("[이미지 없음]")
        return

    if not os.path.exists(image_path):
        doc.add_paragraph(f"[이미지 파일을 찾을 수 없습니다: {image_path}]")
        return

    if not is_supported_docx_image(image_path):
        doc.add_paragraph(f"[Word에 직접 삽입할 수 없는 이미지 형식입니다: {image_path}]")
        return

    try:
        doc.add_picture(image_path, width=Inches(width_inches))
    except Exception as exc:
        doc.add_paragraph(f"[이미지 삽입 실패: {exc}]")


def build_output_docx_path(state: Dict[str, Any]) -> str:
    os.makedirs("outputs", exist_ok=True)
    user_id = sanitize_filename(str(state.get("user_id", "user")))
    consultation_idx = sanitize_filename(str(state.get("consultation_idx", "session")))
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"final_patent_specification_{user_id}_{consultation_idx}_{timestamp}.docx"
    return os.path.abspath(os.path.join("outputs", filename))


def _collect_sections(specification: Dict[str, str], section_groups: List[tuple[str, List[str]]]):
    collected = []
    for title, keys in section_groups:
        content = ""
        for key in keys:
            value = specification.get(key)
            if value and str(value).strip():
                content = str(value).strip()
                break
        if content:
            collected.append((title, content))
    return collected


def add_specification_sections(doc: Document, spec_sections: Dict[str, str]) -> None:
    for title, content in _collect_sections(spec_sections, SECTION_MAPPING[:2]):
        doc.add_paragraph(title)
        doc.add_paragraph(content)

    invention_content_sections = _collect_sections(spec_sections, SECTION_MAPPING[2:5])
    if invention_content_sections:
        doc.add_paragraph("【발명(고안)의 내용】")
        for title, content in invention_content_sections:
            doc.add_paragraph(title)
            doc.add_paragraph(content)

    for title, content in _collect_sections(spec_sections, SECTION_MAPPING[5:]):
        doc.add_paragraph(title)
        doc.add_paragraph(content)


def _claim_sort_key(claim: Any) -> int:
    if isinstance(claim, dict):
        claim_no = claim.get("claim_no")
    else:
        claim_no = None

    if isinstance(claim_no, int):
        return claim_no

    if isinstance(claim_no, str):
        try:
            return int(claim_no)
        except ValueError:
            return 10**9

    return 10**9


def build_claims_text_from_claims(claims: Any) -> str:
    if not isinstance(claims, list):
        return ""

    normalized_claims = []
    for index, claim in enumerate(claims):
        if isinstance(claim, dict):
            text = claim.get("text") or claim.get("claim_text") or claim.get("content") or ""
            claim_no = claim.get("claim_no", index + 1)
        else:
            text = str(claim)
            claim_no = index + 1

        normalized_claims.append((claim_no, str(text).strip()))

    normalized_claims.sort(key=lambda item: (_claim_sort_key({"claim_no": item[0]}), item[0]))

    parts = [f"【청구항 {claim_no}】\n{text}" for claim_no, text in normalized_claims]
    return "\n\n".join(parts)


def add_claims_to_docx(doc: Document, claims: Any) -> None:
    if not isinstance(claims, list):
        return

    normalized_claims = []
    for index, claim in enumerate(claims):
        if isinstance(claim, dict):
            text = claim.get("text") or claim.get("claim_text") or claim.get("content") or ""
            claim_no = claim.get("claim_no", index + 1)
        else:
            text = str(claim)
            claim_no = index + 1

        normalized_claims.append((claim_no, str(text).strip()))

    normalized_claims.sort(key=lambda item: (_claim_sort_key({"claim_no": item[0]}), item[0]))

    for claim_no, text in normalized_claims:
        doc.add_paragraph(f"【청구항 {claim_no}】")
        doc.add_paragraph(text)


def _drawings_to_markdown(drawings: List[Dict[str, Any]]) -> str:
    if not drawings:
        return "도면 없음"

    parts = []
    for index, drawing in enumerate(drawings, start=1):
        fig_no = drawing.get("figure_no", drawing.get("fig_no", index))
        if isinstance(fig_no, int) or str(fig_no).isdigit():
            fig_label = f"【도 {fig_no}】"
        elif str(fig_no).startswith("도 "):
            fig_label = f"【{fig_no}】"
        else:
            fig_label = f"【도 {fig_no}】"

        img_path = get_drawing_image_path(drawing)
        lines = [f"### {fig_label}"]
        lines.append(f"이미지: {img_path}" if img_path else "이미지 없음")
        parts.append("\n".join(lines))

    return "\n\n".join(parts)


def _specification_to_markdown(specification: Dict[str, Any]) -> str:
    sections = []
    for title, content in _collect_sections(specification, SECTION_MAPPING[:2]):
        sections.append(f"### {title}\n{content}")

    invention_content_sections = _collect_sections(specification, SECTION_MAPPING[2:5])
    if invention_content_sections:
        sections.append("### 【발명(고안)의 내용】")
        for title, content in invention_content_sections:
            sections.append(f"#### {title}\n{content}")

    for title, content in _collect_sections(specification, SECTION_MAPPING[5:]):
        sections.append(f"### {title}\n{content}")

    return "\n\n".join(sections)


def render_markdown(
    title: str,
    abstract_text: str,
    representative_drawing_path: Optional[str],
    claims_text: str,
    specification: Dict[str, Any],
    drawings: List[Dict[str, Any]],
) -> str:
    markdown = [f"# {title}", "", "## 요약", abstract_text.strip(), "", "## 대표도"]
    if representative_drawing_path:
        markdown.append(f"![대표도]({representative_drawing_path})")
    else:
        markdown.append("대표도 없음")

    markdown.extend(["", "## 청구항", claims_text.strip(), "", "## 발명의 설명", _specification_to_markdown(specification), "", "## 도면", _drawings_to_markdown(drawings)])
    return "\n".join(markdown).strip()


def create_final_docx(
    output_path: str,
    abstract_text: str,
    representative_drawing_path: Optional[str],
    claims: Any,
    claims_text: str,
    spec_sections: Dict[str, str],
    drawings: List[Dict[str, Any]],
) -> None:
    doc = Document()

    doc.add_heading("요약", level=1)
    doc.add_paragraph(abstract_text.strip())
    doc.add_page_break()

    doc.add_heading("대표도", level=1)
    if representative_drawing_path:
        add_image_or_placeholder(doc, representative_drawing_path)
    else:
        doc.add_paragraph("대표도 없음")
    doc.add_page_break()

    doc.add_heading("청구항", level=1)
    if "【청구범위】" not in claims_text:
        doc.add_paragraph("【청구범위】")

    if isinstance(claims, list):
        add_claims_to_docx(doc, claims)
    else:
        doc.add_paragraph(claims_text)

    doc.add_page_break()

    doc.add_heading("발명의 설명", level=1)
    add_specification_sections(doc, spec_sections)
    doc.add_page_break()

    doc.add_heading("도면", level=1)
    for index, drawing in enumerate(drawings):
        fig_no = drawing.get("figure_no", drawing.get("fig_no", index + 1))
        if isinstance(fig_no, int) or str(fig_no).isdigit():
            fig_label = f"【도 {fig_no}】"
        elif str(fig_no).startswith("도 "):
            fig_label = f"【{fig_no}】"
        else:
            fig_label = f"【도 {fig_no}】"

        doc.add_paragraph(fig_label)

        img_path = get_drawing_image_path(drawing)
        add_image_or_placeholder(doc, img_path)
        doc.add_paragraph("")

    doc.save(output_path)
